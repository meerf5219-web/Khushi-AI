from __future__ import annotations

import csv
import logging
import os
from dataclasses import dataclass
from datetime import datetime
from typing import Dict, Iterator, List, Optional, Tuple

from config.rag import DOCUMENTS_DIR

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class LoadedDocument:
    """Represents an extracted document (may contain multiple pages)."""

    collection: str
    filename: str
    path: str
    text_by_page: Dict[int, str]  # page_number -> extracted text
    last_modified: float


class DocumentLoader:
    """Load documents from collection folders.

    Expected layout:
      knowledge/documents/<Collection>/<files...>
    """

    def __init__(self, *, documents_dir: str = DOCUMENTS_DIR) -> None:
        self._documents_dir = documents_dir

    def iter_documents(self) -> Iterator[LoadedDocument]:
        """Yield all supported documents under documents_dir (recursively per collection)."""
        base = self._documents_dir
        if not os.path.isdir(base):
            return

        for collection in os.listdir(base):
            collection_dir = os.path.join(base, collection)
            if not os.path.isdir(collection_dir):
                continue

            for root, _, files in os.walk(collection_dir):
                for filename in files:
                    path = os.path.join(root, filename)
                    ext = os.path.splitext(filename)[1].lower().lstrip(".")

                    try:
                        yield self._load_single_document(
                            collection=collection, filename=filename, path=path, ext=ext
                        )
                    except Exception as exc:  # pragma: no cover
                        logger.warning("Document load failed: %s (%s)", path, exc)

    def _load_single_document(self, *, collection: str, filename: str, path: str, ext: str) -> LoadedDocument:
        last_modified = os.path.getmtime(path)
        text_by_page: Dict[int, str] = {}

        if ext in {"txt"}:
            text_by_page[1] = self._read_text(path)

        elif ext in {"md", "markdown"}:
            text_by_page[1] = self._read_text(path)

        elif ext in {"csv"}:
            text_by_page[1] = self._read_csv_as_text(path)

        elif ext in {"pdf"}:
            text_by_page = self._read_pdf(path)

        elif ext in {"docx", "doc"}:
            text_by_page = self._read_docx(path)

        else:
            # Unsupported file type: yield empty so indexer can skip.
            text_by_page = {}

        return LoadedDocument(
            collection=collection,
            filename=filename,
            path=path,
            text_by_page=text_by_page,
            last_modified=last_modified,
        )

    def _read_text(self, path: str) -> str:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _read_csv_as_text(self, path: str) -> str:
        lines: List[str] = []
        with open(path, "r", encoding="utf-8", errors="ignore", newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                row_text = " | ".join(cell.strip() for cell in row if cell is not None)
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)

    def _read_pdf(self, path: str) -> Dict[int, str]:
        try:
            import pypdf  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise RuntimeError("pypdf is required to load PDFs.") from exc

        reader = pypdf.PdfReader(path)
        pages: Dict[int, str] = {}
        for i, page in enumerate(reader.pages, start=1):
            pages[i] = page.extract_text() or ""
        return pages

    def _read_docx(self, path: str) -> Dict[int, str]:
        try:
            import docx  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise RuntimeError("python-docx is required to load DOCX.") from exc

        d = docx.Document(path)
        # DOCX has no "pages" easily; treat as one page worth of text.
        text = "\n".join(p.text for p in d.paragraphs if p.text)
        return {1: text}
